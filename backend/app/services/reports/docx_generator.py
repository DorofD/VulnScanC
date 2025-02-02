import json
import ast
from io import BytesIO
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from app.repository.queries.snapshots import get_all_snapshot_data
from app.services.svacer.api import Svacer
from app.services.dependency_track.api import DT


def create_svacer_report(project_name, names=[], ids=[]):
    """принимает либо список имен [project_name, branch_name, snapshot_name], либо список id [project_id, branch_id, snapshot_id]"""

    if names:
        project_name = names[0]
        branch_name = names[1]
        snapshot_name = names[2]
    elif ids:
        project_id = ids[0]
        branch_id = ids[1]
        snapshot_id = ids[2]
    else:
        raise Exception(
            'Svacer report error, incorrect data in docx generator')

    def sort_by_severity(data):
        severity_order = {
            'Critical': 1,
            'Major': 2,
            'Normal': 3,
            'Minor': 4,
            'Undefined': 5
        }

        def get_severity_order(item):
            return severity_order.get(item['checker_severity'], float('inf'))
        sorted_data = sorted(data, key=get_severity_order)
        return sorted_data

    svacer = Svacer()

    if ids:
        markers = svacer.export_snapshot_markers_by_ids(
            project_id, branch_id, snapshot_id)
    elif names:
        markers = svacer.export_snapshot_markers_by_names(
            project_name, branch_name, snapshot_name)

    input_data = markers
    input_data = sort_by_severity(input_data)

    doc = Document()
    heading = doc.add_paragraph(project_name)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.bold = True

    for section in doc.sections:
        # отступы от краев страницы
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)

    def create_single_table(doc, input_data):
        header_values = ['№', 'Описание уязвимости',
                         'Путь до файла', 'Номер строки', 'Комментарий', 'Severity']

        table = doc.add_table(rows=1, cols=0)
        # изменение размеров ячейки по содержимому - False
        table.autofit = False
        table.add_column(Cm(1))
        table.add_column(Cm(5))
        table.add_column(Cm(5))
        table.add_column(Cm(2))
        table.add_column(Cm(4))
        table.add_column(Cm(2.3))
        header_cells = table.rows[0].cells

        for i in range(len(header_values)):
            header_cells[i].text = header_values[i]

        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        count = 1
        for note in input_data:
            table.add_row()

            body_cells = table.rows[count].cells
            body_cells[0].text = str(count)
            body_cells[1].text = f"{note['warnClass']}\n\n{note['msg']}"
            body_cells[2].text = note['file']
            body_cells[3].text = str(note['line'])
            comments = ''
            if 'comments' in note:
                for comment in note['comments']:
                    comments += f"{comment['createdBy']}:\n{comment['text']}\n\n"
            body_cells[4].text = comments
            body_cells[5].text = str(note['checker_severity'])
            count += 1
        tbl = table._tbl
        tblBorders = parse_xml(r'<w:tblBorders %s>'
                               r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'</w:tblBorders>' % nsdecls('w'))
        tbl.tblPr.append(tblBorders)
    create_single_table(doc, input_data)
    report_name = f"{project_name}.docx".replace(
        ':', '-').replace(' ', '_')
    # doc.save(report_name)
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return {'report_name': report_name, 'report': doc_io}


def create_osv_report(snapshot_id: int, severities_string: str):
    snapshot_data = get_all_snapshot_data(snapshot_id)
    project_name = snapshot_data['project_name']
    datetime = snapshot_data['datetime']

    doc = Document()
    heading = doc.add_paragraph(f"{project_name} от {datetime}")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.bold = True
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.size = Pt(16)
    for section in doc.sections:
        # отступы от краев страницы
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)

    def create_component_table(doc, component, component_number):
        header_component_values = [
            '№', 'Информация о компоненте']
        last_paragraph = doc.add_paragraph()
        last_paragraph.paragraph_format.space_after = Cm(1)

        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(0.2)
        table.columns[1].width = Cm(6)
        header_cells = table.rows[0].cells
        for i in range(len(header_component_values)):
            header_cells[i].text = header_component_values[i]
        header_cells[0].width = Cm(1)
        header_cells[1].width = Cm(18)

        # настройка шрифта
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.add_row()
        body_cells = table.rows[1].cells
        body_cells[0].text = str(component_number)
        body_cells[1].text = f"""Путь в проекте: {component['path']}
                                \nТип: {component['type']}
                                \nСсылка: {component['address']}
                                \nТег: {component['tag']}
                                \nВерсия: {component['version']}"""
        # настройка шрифта
        for cell in body_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
        # непрозрачные границы таблицы
        tbl = table._tbl
        tblBorders = parse_xml(r'<w:tblBorders %s>'
                               r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'</w:tblBorders>' % nsdecls('w'))
        tbl.tblPr.append(tblBorders)

    def create_vulnerabilities_table(doc, vulnerabilities):
        header_component_values = [
            'ID', 'Информация об уязвимости', 'Ссылки', 'Severity']
        last_paragraph = doc.add_paragraph()
        last_paragraph.paragraph_format.space_after = Cm(1)
        table = doc.add_table(rows=1, cols=4)
        table.autofit = False

        header_cells = table.rows[0].cells
        for i in range(len(header_component_values)):
            header_cells[i].text = header_component_values[i]
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        count = 1
        for vuln in vulnerabilities:
            # dict_vuln = ast.literal_eval(vuln)
            dict_vuln = vuln
            table.add_row()
            body_cells = table.rows[count].cells
            body_cells[0].text = dict_vuln['id']
            body_cells[1].text = dict_vuln['details']

            ref_str = ''
            for reference in dict_vuln['references']:
                ref_str += f"Тип: {reference['type']}\nСсылка: {reference['url']} \n\n"
            body_cells[2].text = ref_str
            if 'severity' in dict_vuln:
                body_cells[3].text = f"""Тип: {dict_vuln['severity'][0]['type']}
                                        Score: {dict_vuln['severity'][0]['score']}
                                        Base severity: {dict_vuln['severity'][0]['calculated_severities']['base_severity']}
                                        Temporal severity: {dict_vuln['severity'][0]['calculated_severities']['temporal_severity']}
                                        Environmental severity: {dict_vuln['severity'][0]['calculated_severities']['environmental_severity']}
                                    """
            else:
                body_cells[3].text = "No severity found"

            # настройка шрифта
            for cell in body_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
            count += 1
            # непрозрачные границы таблицы
        tbl = table._tbl
        tblBorders = parse_xml(r'<w:tblBorders %s>'
                               r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'</w:tblBorders>' % nsdecls('w'))
        tbl.tblPr.append(tblBorders)

    def sort_vulnerabilities(vulnerabilities, severities_str):
        critical_list = []
        high_list = []
        medium_list = []
        low_list = []
        not_found_list = []

        for vuln in vulnerabilities:
            vuln = ast.literal_eval(vuln)
            if 'severity' not in vuln:
                not_found_list.append(vuln)
                continue
            if vuln['severity'][0]['calculated_severities']['base_severity'] == 'Critical' or vuln['severity'][0]['calculated_severities']['environmental_severity'] == 'Critical' or vuln['severity'][0]['calculated_severities']['temporal_severity'] == 'Critical':
                critical_list.append(vuln)
                continue
            if vuln['severity'][0]['calculated_severities']['base_severity'] == 'High' or vuln['severity'][0]['calculated_severities']['environmental_severity'] == 'High' or vuln['severity'][0]['calculated_severities']['temporal_severity'] == 'High':
                high_list.append(vuln)
                continue
            if vuln['severity'][0]['calculated_severities']['base_severity'] == 'Medium' or vuln['severity'][0]['calculated_severities']['environmental_severity'] == 'Medium' or vuln['severity'][0]['calculated_severities']['temporal_severity'] == 'Medium':
                medium_list.append(vuln)
                continue
            if vuln['severity'][0]['calculated_severities']['base_severity'] == 'Low' or vuln['severity'][0]['calculated_severities']['environmental_severity'] == 'Low' or vuln['severity'][0]['calculated_severities']['temporal_severity'] == 'Low':
                low_list.append(vuln)
                continue

        result = []
        if 'Critical' in severities_str:
            result.extend(critical_list)
        if 'High' in severities_str:
            result.extend(high_list)
        if 'Medium' in severities_str:
            result.extend(medium_list)
        if 'Low' in severities_str:
            result.extend(low_list)
        if 'Without' in severities_str:
            result.extend(not_found_list)

        return result

    count = 1
    components = snapshot_data['components']

    for component in components:
        if 'vulnerabilities' not in component or component['status'] != 'confirmed':
            continue
        sorted_vulnerabilities = sort_vulnerabilities(
            component['vulnerabilities'], severities_string)
        component['vulnerabilities'] = sorted_vulnerabilities
        if 'vulnerabilities' in component and component['vulnerabilities'] != []:
            create_component_table(doc, component, count)
            create_vulnerabilities_table(doc, component['vulnerabilities'])
        count += 1

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return {'project_name': project_name, 'datetime': datetime, 'report': doc_io}


def create_dependency_track_report(project_uuid: str):
    dt = DT()
    components = dt.get_vulnerable_project_components(project_uuid)
    project_name = dt.get_project_info(project_uuid)['name']

    doc = Document()
    heading = doc.add_paragraph(f"{project_name}")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.bold = True
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.size = Pt(16)
    for section in doc.sections:
        # отступы от краев страницы
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)

    def create_component_table(doc, component, component_number):
        header_component_values = [
            '№', 'Информация о компоненте']
        last_paragraph = doc.add_paragraph()
        last_paragraph.paragraph_format.space_after = Cm(1)

        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(0.2)
        table.columns[1].width = Cm(6)
        header_cells = table.rows[0].cells
        for i in range(len(header_component_values)):
            header_cells[i].text = header_component_values[i]
        header_cells[0].width = Cm(1)
        header_cells[1].width = Cm(18)

        # настройка шрифта
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.add_row()
        body_cells = table.rows[1].cells
        body_cells[0].text = str(component_number)
        body_cells[1].text = f"""Название: {component['name']}
                                \nТип: {component['classifier']}
                                \nОписание: {component['description']}
                                \nВерсия: {component['version']}"""
        # настройка шрифта
        for cell in body_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
        # непрозрачные границы таблицы
        tbl = table._tbl
        tblBorders = parse_xml(r'<w:tblBorders %s>'
                               r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'</w:tblBorders>' % nsdecls('w'))
        tbl.tblPr.append(tblBorders)

    def create_vulnerabilities_table(doc, vulnerabilities):
        header_component_values = [
            'CVE id', 'Информация об уязвимости', 'Ссылки', 'Severity']
        last_paragraph = doc.add_paragraph()
        last_paragraph.paragraph_format.space_after = Cm(1)
        table = doc.add_table(rows=1, cols=4)
        table.autofit = False

        header_cells = table.rows[0].cells
        for i in range(len(header_component_values)):
            header_cells[i].text = header_component_values[i]
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        count = 1
        for vuln in vulnerabilities:
            table.add_row()
            body_cells = table.rows[count].cells
            body_cells[0].text = vuln['vulnId']
            body_cells[1].text = vuln['description']
            body_cells[2].text = vuln['references']
            body_cells[3].text = vuln['severity']

            # настройка шрифта
            for cell in body_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
            count += 1
            # непрозрачные границы таблицы
        tbl = table._tbl
        tblBorders = parse_xml(r'<w:tblBorders %s>'
                               r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                               r'</w:tblBorders>' % nsdecls('w'))
        tbl.tblPr.append(tblBorders)

    count = 1

    for component in components:
        create_component_table(doc, component, count)
        create_vulnerabilities_table(doc, component['vulnerabilities'])
        count += 1

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return {'report_name': f"{project_name}.docx", 'report': doc_io}
